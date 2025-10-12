from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
import numpy as np
from io import BytesIO
from scipy.io import wavfile

fsize=[2**8,2**12,2**16]


def plotAudio(Signal, Fs, axs, fsize_value, TimeMargin=[0, 0.02]):
    Signal_cut = Signal[:fsize_value] if len(Signal) >= fsize_value else Signal
    
    N = len(Signal_cut)
    time = np.arange(N) / Fs
    
    fft_signal = np.fft.fft(Signal_cut)
    
    N_half = N // 2
    fft_half = fft_signal[:N_half]
    
    freqs = np.fft.fftfreq(N, 1/Fs)[:N_half]
    
    magnitude = np.abs(fft_half)
    magnitude_db = 20 * np.log10(magnitude + 1e-10)
    
    max_idx = np.argmax(magnitude_db)
    max_freq = freqs[max_idx]
    max_amplitude = magnitude_db[max_idx]
    
    axs[0].plot(time, Signal_cut, linewidth=0.5)
    axs[0].set_xlabel('Czas [s]')
    axs[0].set_ylabel('Amplituda')
    axs[0].set_title('Sygnał audio w dziedzinie czasu')
    axs[0].grid(True, alpha=0.3)
    axs[0].set_xlim(TimeMargin)
    
    axs[1].plot(freqs, magnitude_db, linewidth=0.5)
    axs[1].plot(max_freq, max_amplitude, 'ro', markersize=8, label=f'Max: {max_freq:.2f} Hz')
    axs[1].set_xlabel('Częstotliwość [Hz]')
    axs[1].set_ylabel('Amplituda [dB]')
    axs[1].set_title('Widmo sygnału (połówka)')
    axs[1].grid(True, alpha=0.3)
    axs[1].legend()
    
    return max_freq, max_amplitude


document = Document()
document.add_heading('Analiza plików audio z różnymi rozmiarami FFT', 0)

files = ['sound1.wav']
TimeMargin = [0, 0.02] 

for file in files:
    document.add_heading(f'Plik - {file}', 2)
    
    try:
        Fs, signal = wavfile.read(file)
        
        if len(signal.shape) > 1:
            signal = signal[:, 0]
        
        signal = signal.astype(float)
        
        document.add_paragraph(f'Częstotliwość próbkowania: {Fs} Hz')
        document.add_paragraph(f'Długość sygnału: {len(signal)} próbek ({len(signal)/Fs:.3f} s)')
        document.add_paragraph('')
        
    except FileNotFoundError:
        document.add_paragraph(f'BŁĄD: Nie znaleziono pliku {file}')
        continue
    except Exception as e:
        document.add_paragraph(f'BŁĄD przy wczytywaniu pliku: {str(e)}')
        continue
    
    for fs in fsize:
        document.add_heading(f'FFT size = {fs} ({fs} próbek)', 3)
        
        fig, axs = plt.subplots(2, 1, figsize=(10, 7))
        
        max_freq, max_amplitude = plotAudio(signal, Fs, axs, fs, TimeMargin)
        
        fig.suptitle(f'Analiza pliku {file} - FFT size = {fs}')
        fig.tight_layout(pad=1.5)
        
        memfile = BytesIO()
        fig.savefig(memfile, dpi=150)
        plt.close(fig)
        
        document.add_picture(memfile, width=Inches(6))
        memfile.close()
        
        document.add_paragraph(f'Rozmiar FFT: {fs} próbek')
        document.add_paragraph(f'Częstotliwość maksimum widma: {max_freq:.2f} Hz')
        document.add_paragraph(f'Amplituda przy maksimum: {max_amplitude:.2f} dB')
        document.add_paragraph('')

document.save('report.docx')
print('Raport został zapisany do pliku report.docx')
